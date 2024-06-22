import os
import argparse
import enum_site
import exploit_site
import filetrimmer
import time
from docx import Document

def get_args():
    parser = argparse.ArgumentParser() # create our instance

    parser.add_argument("-u", "--url", dest="url", help="URL") # could be explored in later versions to explore expanding this to include proxies etc?

    options = parser.parse_args()
    if not options.url: # if we dont get the args specified above
        parser.print_help() 
        exit()

    return options

def enumeration(url):

    print ("[+] Performing enumeration of " + url)
    nmap = enum_site.run_nmap(url)
    print("[-] Here is an overview of the IP address provided: \n" + nmap)
    time.sleep(5) # give the user a chance to see what the output is whilst the remainder of the script runs
    dirb = enum_site.run_dirb(url)
    print("\n[-] Here is a list of directories and files on " + url)
    for i in range(1000): # possibility of 65535 open ports so run through that many times
        for resp_code, filepath in dirb.items():
            try:
                if len(filepath[i]) == 0 or resp_code in range(402,599): # 401 is important because it means that the resource exists but requires login, everything else falls under Client/Server errors
                    continue
                else:
                    f1 = open("dirb.txt", "a") # write everything that is <=401
                    f1.write(str(resp_code) + " response to http://" + url + "/" + str(filepath[i] + "\n"))
                    f1.close()
                    f2 = open("dirbstripped.txt", "a") # write the filepath, used in LFI checking
                    f2.write(str(filepath[i] + "\n"))
                    f2.close()
                    print("[-] " + str(resp_code) + " response to http://" + url + "/" + str(filepath[i]))
            except IndexError:
                continue
    return nmap

def exploitation(url):

    xss = exploit_site.xss_runner(url)
    exploit_site.lfi_scan(url)
    return xss

def document_data(url, nmap, xss):
    print("[+] Creating your document")
    document = Document()
    document.add_heading("Penetration Test Report against " + url, level=0)
    document.add_heading("Enumeration", level=1)
    document.add_heading("Nmap", level=2)
    document.add_paragraph(nmap)
    document.add_heading("Dirbuster", level=2)
    f = open("dirb.txt", "r") # instead of transferring this across functions, just read from the file, easier instead of parsing dict entries and handling them
    dirb = f.read()
    f.close()  
    document.add_paragraph(str(dirb))
    document.add_heading("Exploitation", level=1)
    document.add_heading("XSS", level=2)
    count = 0
    while True:
        try:
            document.add_paragraph(xss[count])
            count += 1
        except:
            break
    document.add_heading("LFI", level=2)
    document.add_paragraph("WAVARG cannot be certain that LFI is present, although suspects if possible, it may be at: ")
    count = 0
    fn = "results"
    filetrimmer.run()
    while True:
        try:
            count += 1
            filename = "finalfinal_"+fn + str(count) + ".txt" # e.g. finalfinal_results1.txt - messy but does the job
            print ("[+]" + filename)
            lfi_fp = exploit_site.get_lfi_fp() # reading dirbstripped.txt
            file = open(filename, "r") 
            document.add_paragraph("Scan "+str(count)+" on http://" + str(url) + "/"+ lfi_fp[count] + "\n" + str(file.read())) # put it all together and send it
        except:
            break
    document.save("demo.docx")
    print("[+] Document Saved to demo.docx")
    os.system("rm *dirb*.txt *results*") # clean up so we remove hundreds of txt files
def main():
    options = get_args() 
    url = options.url
    # fetch the arguments and give them local variables
    
    nmap = enumeration(url)
    xss = exploitation(url)
    document_data(url, nmap, xss)

    # run through our various functions
    


if __name__ == "__main__":
    main()